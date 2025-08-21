import pandas as pd
from docx import Document
import win32com.client
import re
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from ttkthemes import ThemedTk

class PlaceholderEntry(ttk.Entry):
    def __init__(self, container, placeholder, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        self.placeholder = placeholder
        self.placeholder_color = 'grey'
        self.default_fg_color = self['foreground']
        
        self.bind("<FocusIn>", self._clear_placeholder)
        self.bind("<FocusOut>", self._add_placeholder)
        
        self._add_placeholder()
    
    def _clear_placeholder(self, e):
        if self.get() == self.placeholder:
            self.delete(0, tk.END)
            self['foreground'] = self.default_fg_color
    
    def _add_placeholder(self, e=None):
        if not self.get():
            self.insert(0, self.placeholder)
            self['foreground'] = self.placeholder_color

class NotadorGUI:
    def __init__(self):
        self.root = ThemedTk(theme="arc")
        self.root.title("Notador - Generador de Reportes")
        self.root.geometry("1000x700")
        
        # Configurar el tema y estilos
        style = ttk.Style()
        style.configure('TLabel', font=('Arial', 10))
        style.configure('TButton', font=('Arial', 10), padding=5)
        style.configure('Heading.TLabel', font=('Arial', 12, 'bold'))
        style.configure('TLabelframe.Label', font=('Arial', 10, 'bold'))
        
        # Estilos personalizados para cada sección
        style.configure('Files.TLabelframe', background='#FFE6F3')  # Rosa suave
        style.configure('Files.TLabelframe.Label', font=('Arial', 10, 'bold'), background='#FFE6F3')
        
        style.configure('Students.TLabelframe', background='#E6E6FA')  # Lila suave
        style.configure('Students.TLabelframe.Label', font=('Arial', 10, 'bold'), background='#E6E6FA')
        
        style.configure('Progress.TLabelframe', background='#FFD6FF')  # Fucsia suave
        style.configure('Progress.TLabelframe.Label', font=('Arial', 10, 'bold'), background='#FFD6FF')
        
        # Colores personalizados para los widgets
        self.root.configure(background='#f0f0f0')
        style.configure('Custom.TButton',
                      background='#0078D7',
                      foreground='white',
                      padding=10)
        style.map('Custom.TButton',
                 background=[('active', '#1e88e5')],
                 foreground=[('active', 'white')])
                 
        self.notador = Notador()
        self.current_data = {}  # Almacena los datos del Excel actual
        self.setup_gui()
        
    def setup_gui(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configurar el peso de las columnas y filas
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Configurar pesos de las filas para mejor distribución
        main_frame.rowconfigure(0, weight=0)  # Panel de archivos - altura fija
        main_frame.rowconfigure(1, weight=10)  # Panel de grados/estudiantes - máxima expansión
        main_frame.rowconfigure(2, weight=3)  # Panel de progreso - expansión moderada
        
        # Sección de archivos
        files_frame = ttk.LabelFrame(main_frame, text="Archivos", padding="5", style='Files.TLabelframe')
        files_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N), pady=5)
        
        # Excel
        ttk.Label(files_frame, text="Archivo Excel:").grid(row=0, column=0, sticky=tk.W)
        self.excel_path_var = tk.StringVar()
        ttk.Entry(files_frame, textvariable=self.excel_path_var, width=50).grid(row=0, column=1, padx=5)
        ttk.Button(files_frame, text="Buscar", command=self.browse_excel).grid(row=0, column=2)
        
        # Word
        ttk.Label(files_frame, text="Plantilla Word:").grid(row=1, column=0, sticky=tk.W)
        self.word_path_var = tk.StringVar()
        ttk.Entry(files_frame, textvariable=self.word_path_var, width=50).grid(row=1, column=1, padx=5)
        ttk.Button(files_frame, text="Buscar", command=self.browse_word).grid(row=1, column=2)
        
        # Carpeta de salida
        ttk.Label(files_frame, text="Carpeta de salida:").grid(row=2, column=0, sticky=tk.W)
        self.output_folder_var = tk.StringVar()
        ttk.Entry(files_frame, textvariable=self.output_folder_var, width=50).grid(row=2, column=1, padx=5)
        ttk.Button(files_frame, text="Buscar", command=self.browse_output_folder).grid(row=2, column=2)
        
        # Botón para cargar datos
        ttk.Button(files_frame, text="Cargar Datos", 
                  command=self.load_data).grid(row=3, column=0, columnspan=3, pady=5)
        
        # Panel izquierdo - Lista de grados
        grades_frame = ttk.LabelFrame(main_frame, text="Grados", padding="5", style='Students.TLabelframe')
        grades_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.grades_listbox = tk.Listbox(grades_frame, width=20)
        self.grades_listbox.pack(fill=tk.BOTH, expand=True)
        self.grades_listbox.bind('<<ListboxSelect>>', self.on_grade_select)
        
        # Panel derecho - Lista de estudiantes
        students_frame = ttk.LabelFrame(main_frame, text="Estudiantes", padding="5", style='Students.TLabelframe')
        students_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5, padx=5)
        students_frame.columnconfigure(0, weight=1)  # Hacer que la columna se expanda
        students_frame.rowconfigure(0, weight=0)  # Fila de búsqueda - altura fija
        students_frame.rowconfigure(1, weight=1)  # Fila del Treeview - expansión máxima
        students_frame.rowconfigure(2, weight=0)  # Fila de botones - altura fija
        
        # Barra de búsqueda
        self.search_var = tk.StringVar()
        self.search_var.trace('w', self.filter_students)
        self.search_entry = PlaceholderEntry(students_frame, 
                                           placeholder="Buscar estudiante...",
                                           textvariable=self.search_var)
        self.search_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=5, pady=5)
        
        # Estilo personalizado para el Treeview
        style = ttk.Style()
        style.configure("Custom.Treeview",
                      background="#f0f0f0",
                      foreground="black",
                      rowheight=25,
                      fieldbackground="#f0f0f0")
        style.configure("Custom.Treeview.Heading",
                      background="#e0e0e0",
                      foreground="black",
                      font=('Arial', 10, 'bold'))
        style.map("Custom.Treeview",
                 background=[('selected', '#0078D7')],
                 foreground=[('selected', 'white')])

        # Lista de estudiantes con scrollbar
        student_list_frame = ttk.Frame(students_frame)
        student_list_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
        student_list_frame.columnconfigure(0, weight=1)
        student_list_frame.rowconfigure(0, weight=1)
        
        # Treeview con checkboxes y estilo moderno
        self.students_tree = ttk.Treeview(student_list_frame, 
                                        columns=('check', 'id', 'name', 'group'),
                                        show='headings', 
                                        selectmode='browse',
                                        style="Custom.Treeview")
        self.students_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        scrollbar = ttk.Scrollbar(student_list_frame, orient=tk.VERTICAL, command=self.students_tree.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Configurar columnas
        self.students_tree.heading('check', text='✓')
        self.students_tree.heading('id', text='ID')
        self.students_tree.heading('name', text='Nombre')
        self.students_tree.heading('group', text='Grupo')
        
        self.students_tree.column('check', width=30, anchor='center')
        self.students_tree.column('id', width=100)
        self.students_tree.column('name', width=250)
        self.students_tree.column('group', width=80, anchor='center')
        
        # Variable para controlar el estado de los checkboxes
        self.checked_items = set()
        
        # Evento de click en el árbol
        def on_tree_click(event):
            region = self.students_tree.identify_region(event.x, event.y)
            if region == "cell":
                column = self.students_tree.identify_column(event.x)
                if column == '#1':  # Columna del checkbox
                    item = self.students_tree.identify_row(event.y)
                    if item in self.checked_items:
                        self.checked_items.remove(item)
                        self.students_tree.set(item, 'check', '☐')
                    else:
                        self.checked_items.add(item)
                        self.students_tree.set(item, 'check', '☒')
        
        self.students_tree.bind('<Button-1>', on_tree_click)
        self.students_tree.config(yscrollcommand=scrollbar.set)
        
        # Botones de acción
        button_frame = ttk.Frame(students_frame)
        button_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), padx=5, pady=5)
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)
        
        ttk.Button(button_frame, text="Procesar Seleccionados", 
                  command=self.process_selected).grid(row=0, column=0, padx=(0,2), sticky=(tk.W, tk.E))
        ttk.Button(button_frame, text="Procesar Todos", 
                  command=self.process_all).grid(row=0, column=1, padx=(2,0), sticky=(tk.W, tk.E))
        
        # Panel inferior - Progreso
        progress_frame = ttk.LabelFrame(main_frame, text="Progreso", padding="5", style='Progress.TLabelframe')
        progress_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        progress_frame.columnconfigure(0, weight=1)  # Hacer que la columna se expanda
        progress_frame.rowconfigure(1, weight=1)  # Hacer que la fila del área de texto se expanda
        
        # Frame para la barra de progreso y el porcentaje
        progress_bar_frame = ttk.Frame(progress_frame)
        progress_bar_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Barra de progreso
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_bar_frame, 
            variable=self.progress_var,
            maximum=100,
            mode='determinate',
            length=300
        )
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Etiqueta de porcentaje
        self.progress_label = ttk.Label(progress_bar_frame, text="0%")
        self.progress_label.pack(side=tk.LEFT, padx=5)
        
        # Área de texto con scrollbar
        text_frame = ttk.Frame(progress_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.progress_text = tk.Text(text_frame, height=5, width=70)
        self.progress_text.pack(fill=tk.BOTH, expand=True)
        
        # Configurar el scrollbar
        scrollbar.config(command=self.progress_text.yview)
        self.progress_text.config(yscrollcommand=scrollbar.set)
        
    def browse_excel(self):
        filename = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if filename:
            self.excel_path_var.set(filename)
            
    def browse_word(self):
        filename = filedialog.askopenfilename(
            title="Seleccionar plantilla Word",
            filetypes=[("Word files", "*.docx")]
        )
        if filename:
            self.word_path_var.set(filename)
            
    def browse_output_folder(self):
        folder = filedialog.askdirectory(
            title="Seleccionar carpeta para guardar boletines"
        )
        if folder:
            self.output_folder_var.set(folder)
            self.notador.set_output_folder(folder)
    
    def normalize_column_names(self, df):
        """Normaliza los nombres de las columnas para manejar variaciones"""
        # Primero limpiar los nombres de las columnas
        df.columns = [str(col).strip() for col in df.columns]
        
        column_mapping = {
            'GRUPO': ['GRUPO', 'grupo', 'Group', 'group', 'GRUPO.1', 'Grupo.1'],
            'PERIODO': ['PERIODO', 'periodo', 'PERÍODO', 'período', 'PERIODO.1'],
            'estudiante': ['estudiante', 'ESTUDIANTE', 'Estudiante', 'NOMBRE', 'NOMBRES Y APELLIDOS']
        }
        
        renamed_columns = {}
        for standard_name, variations in column_mapping.items():
            for variant in variations:
                if variant in df.columns:
                    renamed_columns[variant] = standard_name
                    break
        
        if renamed_columns:
            df = df.rename(columns=renamed_columns)
        return df

    def load_data(self):
        try:
            excel_path = self.excel_path_var.get()
            if not excel_path:
                messagebox.showerror("Error", "Por favor seleccione el archivo Excel")
                return
                
            self.notador.set_excel_file(excel_path)
            xlsx = pd.ExcelFile(excel_path)
            
            # Limpiar listas actuales
            self.grades_listbox.delete(0, tk.END)
            self.students_tree.delete(*self.students_tree.get_children())
            self.current_data = {}
            
            # Cargar grados, excluyendo la pestaña consolidado
            for grado in xlsx.sheet_names:
                # Ignorar la pestaña consolidado
                if grado.lower().strip() == 'consolidado':
                    continue
                    
                # Leer el Excel manejando celdas combinadas
                df = pd.read_excel(
                    excel_path,
                    sheet_name=grado,
                    header=0,  # Primera fila como encabezados
                    na_filter=False  # No convertir valores vacíos a NaN
                )
                # Normalizar nombres de columnas
                df = self.normalize_column_names(df)
                
                # Verificar columnas requeridas
                required_columns = ['estudiante']
                missing_columns = [col for col in required_columns if col not in df.columns]
                
                if missing_columns:
                    messagebox.showerror("Error", 
                        f"Faltan columnas requeridas en la hoja {grado}: {', '.join(missing_columns)}\n"
                        "El archivo debe tener al menos una columna para el estudiante.")
                    return
                
                # Solo agregar grados válidos a la lista y al diccionario de datos
                self.grades_listbox.insert(tk.END, grado)
                self.current_data[grado] = df
                
            self.add_progress("Datos cargados exitosamente")
                
            self.add_progress("Datos cargados exitosamente")
            
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def on_grade_select(self, event):
        selection = self.grades_listbox.curselection()
        if not selection:
            return
            
        grado = self.grades_listbox.get(selection[0])
        self.load_students_for_grade(grado)
    
    def load_students_for_grade(self, grado):
        # Limpiar lista actual
        self.students_tree.delete(*self.students_tree.get_children())
        
        if grado not in self.current_data:
            return
            
        df = self.current_data[grado]
        
        # Asegurarse de que la columna estudiante sea string
        df['estudiante'] = df['estudiante'].astype(str)
        
        for _, row in df.iterrows():
            student_field = str(row['estudiante']).strip()
            if student_field.lower() == 'nan' or student_field == '':
                continue
                
            student_id, student_name, _, _ = self.notador.parse_student_info(student_field)
            if student_id:
                # Obtener el grupo del estudiante
                grupo = str(row.get('GRUPO', '')).strip()
                if pd.isna(grupo) or grupo.lower() == 'nan' or grupo == '':
                    grupo = 'N/A'
                    
                item = self.students_tree.insert('', tk.END, values=(
                    '☐',  # Checkbox vacío
                    student_id,
                    student_name,
                    grupo
                ))
    
    def filter_students(self, *args):
        search_term = self.search_var.get().lower()
        selection = self.grades_listbox.curselection()
        if not selection:
            return
            
        grado = self.grades_listbox.get(selection[0])
        self.load_students_for_grade(grado)
        
        if search_term:
            for item in self.students_tree.get_children():
                values = self.students_tree.item(item)['values']
                if not any(search_term in str(value).lower() for value in values):
                    self.students_tree.detach(item)
    
    def process_selected(self):
        # Usar items checked en lugar de selección
        if not self.checked_items:
            messagebox.showinfo("Información", "Por favor seleccione al menos un estudiante usando los checkboxes")
            return
            
        if not self.output_folder_var.get():
            messagebox.showerror("Error", "Por favor seleccione una carpeta de salida para los boletines")
            return
        
        # Limpiar área de progreso
        self.progress_text.delete(1.0, tk.END)
        self.progress_var.set(0)
            
        try:
            word_path = self.word_path_var.get()
            if not word_path:
                messagebox.showerror("Error", "Por favor seleccione la plantilla Word")
                return
                
            self.notador.set_word_template(word_path)
            
            total = len(self.checked_items)
            self.progress_var.set(0)
            
            selection = self.grades_listbox.curselection()
            if not selection:
                messagebox.showerror("Error", "Por favor seleccione un grado")
                return
                
            grado = self.grades_listbox.get(selection[0])
            df = self.current_data[grado]
            
            # Asegurarse de que la columna estudiante sea string
            df['estudiante'] = df['estudiante'].astype(str)
            
            for i, item in enumerate(self.checked_items):
                values = self.students_tree.item(item)['values']
                student_id = str(values[1]).strip()  # ID está en la segunda columna (índice 1)
                student_name = str(values[2]).strip()  # Nombre está en la tercera columna (índice 2)
                grupo = str(values[3]).strip()  # Grupo está en la cuarta columna (índice 3)
                
                try:
                    # Obtener el periodo del estudiante desde el DataFrame
                    student_rows = df[df['estudiante'].apply(lambda x: student_id in str(x))]
                    if not student_rows.empty:
                        student_row = student_rows.iloc[0]
                        periodo = str(student_row.get('PERIODO', '')).strip()
                    else:
                        periodo = ''
                    
                    self.add_progress(f"\n📝 Generando boletín para {student_name}")
                    
                    # Actualizar barra de progreso antes de procesar
                    progress = (i / total) * 100
                    self.progress_var.set(progress)
                    self.progress_label.config(text=f"{int(progress)}%")
                    self.root.update()
                    
                    # Procesar el boletín
                    self.notador.process_student(student_id, grado, grupo, periodo, callback=self.add_progress)
                    
                    # Actualizar barra de progreso después de procesar
                    progress = ((i + 1) / total) * 100
                    self.progress_var.set(progress)
                    self.progress_label.config(text=f"{int(progress)}%")
                    self.root.update()
                    
                except Exception as e:
                    self.add_progress(f"❌ Error con {student_name}: {str(e)}")
                    continue
                
            messagebox.showinfo("Éxito", "Procesamiento completado")
            
        except Exception as e:
            messagebox.showerror("Error", str(e))
            
    def add_progress(self, message):
        """Actualiza el área de progreso y la barra de progreso"""
        # Extraer el porcentaje del mensaje si existe
        import re
        percentage_match = re.search(r'\((\d+)%\)', message)
        if percentage_match:
            percentage = int(percentage_match.group(1))
            self.progress_var.set(percentage)
            self.progress_label.config(text=f"{percentage}%")
        
        # Agregar mensaje al área de texto (sin el porcentaje)
        clean_message = re.sub(r'\(\d+%\)', '', message).strip()
        self.progress_text.insert(tk.END, clean_message + "\n")
        self.progress_text.see(tk.END)
        
        # Forzar actualización de la interfaz
        self.root.update()
        
    def process_all(self):
        """Procesa todos los estudiantes del grado seleccionado"""
        try:
            # Verificar que tengamos todo lo necesario
            word_path = self.word_path_var.get()
            if not word_path:
                messagebox.showerror("Error", "Por favor seleccione la plantilla Word")
                return
                
            if not self.output_folder_var.get():
                messagebox.showerror("Error", "Por favor seleccione una carpeta de salida para los boletines")
                return
            
            selection = self.grades_listbox.curselection()
            if not selection:
                messagebox.showerror("Error", "Por favor seleccione un grado")
                return
                
            grado = self.grades_listbox.get(selection[0])
            if grado not in self.current_data:
                messagebox.showerror("Error", "Por favor cargue los datos primero")
                return
            
            # Configurar notador
            self.notador.set_word_template(word_path)
            self.notador.set_output_folder(self.output_folder_var.get())
            
            # Limpiar área de progreso
            self.progress_text.delete(1.0, tk.END)
            self.progress_var.set(0)
            self.progress_label.config(text="0%")
            # Iniciar procesamiento
            df = self.current_data[grado]
            self.add_progress(f"🎯 Iniciando procesamiento del grado {grado}")
            
            # Filtrar filas válidas
            df = df[df['estudiante'].apply(lambda x: str(x).strip().lower() != 'nan' and str(x).strip() != '')]
            total = len(df)
            
            if total == 0:
                messagebox.showwarning("Advertencia", "No hay estudiantes para procesar en este grado")
                return
            
            for i, row in df.iterrows():
                student_field = str(row['estudiante']).strip()
                student_id, student_name, _, _ = self.notador.parse_student_info(student_field)
                
                if student_id:
                    try:
                        # Obtener periodo del estudiante
                        periodo = str(row.get('PERIODO', '')).strip()
                        if pd.isna(periodo) or periodo.lower() == 'nan':
                            periodo = ''
                            
                        # Obtener grupo del estudiante
                        grupo = str(row.get('GRUPO', '')).strip()
                        if pd.isna(grupo) or grupo.lower() == 'nan':
                            grupo = 'N/A'
                        
                        # Procesar estudiante y actualizar progreso
                        self.add_progress(f"\n📝 Procesando boletín para {student_name}")
                        self.progress_var.set((i / total) * 100)
                        self.root.update_idletasks()
                        
                        self.notador.process_student(student_id, grado, grupo, periodo, callback=self.add_progress)
                        
                        # Actualizar progreso final para este estudiante
                        self.progress_var.set(((i + 1) / total) * 100)
                        self.root.update_idletasks()
                        
                    except Exception as e:
                        self.add_progress(f"❌ Error con {student_name}: {str(e)}")
                        continue
            
            # Mostrar mensaje de éxito
            self.add_progress("\n✨ Procesamiento completado exitosamente")
            messagebox.showinfo("Éxito", "Se han generado todos los boletines correctamente")
            
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def run(self):
        self.root.mainloop()

class Notador:
    def __init__(self):
        self.excel_path = None
        self.word_template = None
        self.output_folder = None
        self.debug = False  # Deshabilitar mensajes de debug
        
        # Mapeo de nombres de columnas del Excel
        self.column_mapping = {
            'GRUPO': ['GRUPO', 'grupo', 'Group', 'group', 'GRUPO.1', 'Grupo.1'],
            'PERIODO': ['PERIODO', 'periodo', 'PERÍODO', 'período', 'PERIODO.1'],
            'estudiante': ['estudiante', 'ESTUDIANTE', 'Estudiante', 'NOMBRE', 'NOMBRES Y APELLIDOS']
        }
        
        # Mapeo de áreas entre Excel y Word
        # El texto en el diccionario debe coincidir exactamente con el texto en el Word
        self.areas_mapping = {
            'Investigación': 'Introducción a la investigación Formativa',  # Nombre del área en el Word
            'Ciencias Naturales': 'Ciencias Naturales y Educación Ambiental',
            'Ciencias Sociales': 'Ciencias Sociales, Historia, Geografía, Constitución Política y Democracia',
            'Educ, Artística': 'Educación Artística y Cultural',
            'Edu, Ética': 'Educación Ética y en Valores Humanos',
            'Edu, física': 'Educación Física, Recreación y Deportes',
            'Edu, Religiosa': 'Educación Religiosa',
            'Lengua Castellana': 'Humanidades, Lengua Castellana',
            'Matemáticas': 'Matemáticas',
            'Tecnología e informatica': 'Tecnología e Informática',
            'Inglés': 'Humanidades, idioma extranjero (inglés)'
        }
        
        # Mapeo de campos de Word - usando los textos exactos del documento
        self.word_fields_mapping = {
            'periodo': ['PERÍODO'],  # Solo la versión exacta que está en el documento
            'grado': ['GRADO'],
            'grupo': ['GRUPO'],
            'nombre_completo': ['NOMBRE Y APELLIDOS COMPLETOS DEL ESTUDIANTE:'],  # Con los dos puntos
            'id': ['ID INSTITUCIONAL'],
            'materias_perdidas': ['Número de áreas o asignaturas con nota no aprobatoria'],
            'promedio': ['Promedio Académico']
        }
        
        # Mapeo para las notas de las áreas
        self.word_area_grade_format = '[NOTA DEL PERIODO]'  # Se usará como sufijo para cada área
        
    def normalize_column_names(self, df):
        """Normaliza los nombres de las columnas para manejar variaciones"""
        # Primero limpiar los nombres de las columnas
        df.columns = [str(col).strip() for col in df.columns]
        
        renamed_columns = {}
        for standard_name, variations in self.column_mapping.items():
            for variant in variations:
                if variant in df.columns:
                    renamed_columns[variant] = standard_name
                    break
        
        if renamed_columns:
            df = df.rename(columns=renamed_columns)
        return df
        
    def set_excel_file(self, excel_path):
        """Establece el archivo Excel a usar"""
        if not Path(excel_path).exists():
            raise FileNotFoundError(f"El archivo Excel {excel_path} no existe")
        self.excel_path = excel_path
        
    def set_word_template(self, word_path):
        """Establece la plantilla Word a usar"""
        if not Path(word_path).exists():
            raise FileNotFoundError(f"La plantilla Word {word_path} no existe")
        self.word_template = word_path
        
    def set_output_folder(self, folder_path):
        """Establece la carpeta de salida para los boletines"""
        self.output_folder = Path(folder_path)
    
    def get_excel_areas(self):
        """Obtiene las áreas del archivo Excel"""
        df = pd.read_excel(self.excel_path, sheet_name=0)
        areas = [col for col in df.columns if col not in ['GRADO', 'GRUPO', 'PERIODO', 'estudiante', 'Promedio', 'Mención de honor']]
        return areas

    def get_word_areas(self):
        """Obtiene las áreas de la plantilla Word"""
        doc = Document(self.word_template)
        areas = []
        for paragraph in doc.paragraphs:
            for area_excel, area_word in self.areas_mapping.items():
                if area_word in paragraph.text.upper():
                    areas.append(area_excel)
        return areas
    
    def find_and_replace_in_doc(self, doc, placeholder, value):
        """Busca y reemplaza texto en todo el documento Word"""
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(placeholder, str(value))
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(placeholder, str(value))
                                
    def process_word_document(self, doc_path, pdf_path, replacements, student_row, callback=None, areas_mapping=None):
        """
        Procesa un documento Word, realizando los reemplazos necesarios y llenando las notas.
        Guarda el documento en formato Word y PDF.
        
        Args:
            doc_path (str): Ruta al documento Word
            pdf_path (str): Ruta donde se guardará el PDF
            replacements (dict): Diccionario con los textos a buscar y reemplazar
            student_row (Series): Fila del DataFrame con los datos del estudiante
            callback (callable, optional): Función para reportar progreso
        """
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            doc = word.Documents.Open(doc_path)
            
            # Reemplazar información básica
            if callback:
                callback("Aplicando información básica...")
            
            # Constantes de Word
            wdFindContinue = 1
            wdReplaceAll = 2
            
            # Realizar reemplazos básicos
            for find_text, replace_text in replacements.items():
                # Buscar en todas las tablas del documento
                for table_idx in range(1, doc.Tables.Count + 1):
                    table = doc.Tables(table_idx)
                    for row_idx in range(1, table.Rows.Count + 1):
                        for col_idx in range(1, table.Columns.Count + 1):
                            try:
                                cell = table.Cell(row_idx, col_idx)
                                cell_text = cell.Range.Text.rstrip('\r\x07')
                                
                                # Si encontramos el texto en la celda
                                if find_text.strip() == cell_text.strip():
                                    # Para nombre_completo e id, escribir en la celda de abajo
                                    if find_text in [
                                        'NOMBRE Y APELLIDOS COMPLETOS DEL ESTUDIANTE:',
                                        'ID INSTITUCIONAL'
                                    ]:
                                        if row_idx < table.Rows.Count:  # Verificar que existe una fila siguiente
                                            next_cell = table.Cell(row_idx + 1, col_idx)
                                            next_cell.Range.Text = replace_text
                                    # Para los demás campos, escribir en la celda contigua
                                    else:
                                        if col_idx < table.Columns.Count:  # Verificar que existe una columna siguiente
                                            next_cell = table.Cell(row_idx, col_idx + 1)
                                            next_cell.Range.Text = replace_text
                            except:
                                # Ignorar errores de celdas no existentes
                                continue
            # Buscar la tabla que contiene "ÁREAS"
            table_found = False
            
            for table in doc.Tables:
                if table_found:
                    break
                    
                # Buscar la celda que contiene "ÁREAS"
                for row in table.Rows:
                    for cell in row.Cells:
                        cell_text = cell.Range.Text.strip().rstrip('\r\x07')  # Eliminar caracteres especiales
                        if cell_text == "ÁREAS":
                            current_table = table
                            table_found = True
                            break
                    if table_found:
                        break
            
            if not table_found:
                raise ValueError("No se encontró la tabla de ÁREAS en el documento")
            
            # Procesar las áreas
            areas_procesadas = 0

            # Use the provided mapping if given, otherwise fall back to the default
            mapping_to_use = areas_mapping if areas_mapping is not None else self.areas_mapping

            # Iteramos por las filas de la tabla
            for row in current_table.Rows:
                try:
                    # Obtenemos la celda del área (columna 1) y la celda de la nota (columna 2)
                    area_cell = row.Cells(1)
                    nota_cell = row.Cells(2)
                    
                    area_text = area_cell.Range.Text.strip().rstrip('\r\x07')
                    
                    # Buscar el área correspondiente en el mapeo
                    for area_excel, area_word in mapping_to_use.items():
                        try:
                            if area_word.upper() in area_text.upper():
                                # Si encontramos el área y está en los datos del estudiante
                                if area_excel in student_row:
                                    nota = str(student_row[area_excel]).strip()
                                    nota_cell.Range.Text = nota
                                    areas_procesadas += 1
                                break
                        except Exception:
                            # Ignorar problemas con comparación de textos
                            continue
                except Exception as e:
                    continue
            # Guardar como Word
            doc.Save()
            
            # Guardar como PDF
            if callback:
                callback("💾 Guardando versión PDF...")
            
            # Constantes de Word para PDF
            wdFormatPDF = 17  # Formato PDF
            doc.SaveAs2(pdf_path, FileFormat=wdFormatPDF)
            
            # Cerrar el documento
            doc.Close()
        finally:
            word.Quit()
        
    def map_areas(self):
        """Crea un mapeo entre las áreas de Excel y Word"""
        excel_areas = self.get_excel_areas()
        word_areas = self.get_word_areas()
        
        # Aquí se implementará la lógica para mapear áreas similares
        # Por ejemplo: "Ed. Fisica" -> "Educacion Fisica Recreacion y Deportes"
        # Se podría usar fuzzy matching o reglas predefinidas
        
    def calculate_academic_stats(self, student_row):
        """Calcula estadísticas académicas: promedio y número de materias perdidas"""
        # Set para debug - guardar qué columnas se usaron
        used_columns = []
        
        # Columnas administrativas a excluir (incluyendo variantes comunes)
        admin_columns = {
            'estudiante', 'ESTUDIANTE', 
            'GRUPO', 'grupo',
            'PERIODO', 'periodo', 'PERÍODO',
            'Promedio', 'PROMEDIO', 'promedio',
            'Mención de honor', 'MENCION DE HONOR',
            'OBSERVACIONES', 'observaciones',
            'GRADO', 'grado',
            'Areas deficitadas', 'AREAS DEFICITADAS', 'Áreas deficitadas',
            'areas deficitadas', 'Areas perdidas', 'AREAS PERDIDAS'
        }
        
        # Función para normalizar nombres
        import unicodedata
        def normalize_name(s):
            if s is None:
                return ''
            s = str(s).strip().lower()
            s = unicodedata.normalize('NFKD', s)
            s = ''.join(ch for ch in s if not unicodedata.combining(ch))
            return s
        
        normalized_admin = {normalize_name(c) for c in admin_columns}
        
        notas = []
        materias_perdidas = 0
        
        # Debug - imprimir todas las columnas
        print("\nColumnas en student_row:")
        for column, value in student_row.items():
            print(f"Columna original: '{column}' -> Valor: '{value}'")
        
        print("\nProcesando notas:")
        # Iterar y recolectar notas
        for column, value in student_row.items():
            col_norm = normalize_name(column)
            if not col_norm or col_norm in normalized_admin:
                print(f"Ignorando columna admin: '{column}' (normalizada: '{col_norm}')")
                continue
                
            try:
                nota_str = str(value).strip().replace(',', '.')
                if nota_str and nota_str.lower() != 'nan':
                    nota = float(nota_str)
                    print(f"Añadiendo nota {nota} de columna '{column}'")
                    notas.append(nota)
                    used_columns.append(column)
                    if nota < 3.5:
                        materias_perdidas += 1
            except (ValueError, TypeError):
                print(f"Error convirtiendo valor '{value}' de columna '{column}'")
                continue
        
        # Debug - mostrar resumen
        print(f"\nResumen del cálculo:")
        print(f"Columnas usadas ({len(used_columns)}): {', '.join(used_columns)}")
        print(f"Total notas válidas: {len(notas)}")
        print(f"Notas: {notas}")
        if notas:
            print(f"Suma: {sum(notas)}")
            print(f"Promedio: {sum(notas) / len(notas)}")
        
       
        promedio = round(sum(notas) / len(notas), 2) if notas else 0
        
        return promedio, materias_perdidas
        
    def parse_student_info(self, student_field):
        """Extrae la información del estudiante del campo compuesto"""
        if pd.isna(student_field) or not isinstance(student_field, str):
            return None, None, None, None
            
        # Patrón específico para ID de 9 dígitos seguido de nombre
        pattern = r"(\d{9})\s*-\s*(.+)"
        match = re.match(pattern, student_field)
        if match:
            id_number = match.group(1)
            full_name = match.group(2).strip()
            # Separar nombres y apellidos (asumiendo formato: APELLIDOS NOMBRES)
            parts = full_name.split()
            if len(parts) >= 2:
                # Asumimos que los dos primeros son apellidos
                apellidos = " ".join(parts[:2])
                nombres = " ".join(parts[2:])
            else:
                apellidos = full_name
                nombres = ""
            return id_number, full_name, apellidos, nombres
        return None, None, None, None
        
    def process_student(self, student_id, grado, grupo, periodo=None, callback=None):
        """Procesa la información de un estudiante y genera su documento"""
        try:
            # 20% - Cargar datos del Excel
            if callback:
                callback("⌛ Cargando datos del estudiante (20%)")
            df = pd.read_excel(
                self.excel_path,
                sheet_name=str(grado),
                header=0,
                na_filter=False
            )
            df = self.normalize_column_names(df)
            
            # Asegurarse de que todos los datos sean strings
            for col in df.columns:
                df[col] = df[col].astype(str).str.strip()
            
            # 40% - Buscar estudiante
            if callback:
                callback("🔍 Localizando información del estudiante (40%)")
            student_rows = df[df['estudiante'].apply(lambda x: str(student_id) in str(x))]
            
            if student_rows.empty:
                raise ValueError(f"No se encontró el estudiante con ID {student_id}")
                
            student_row = student_rows.iloc[0]
                
            # 60% - Preparar información
            if callback:
                callback("📋 Preparando la información (60%)")
            _, nombre_completo, apellidos, nombres = self.parse_student_info(student_row['estudiante'])
            
            # Verificar que existan los recursos necesarios
            if not self.word_template or not Path(self.word_template).exists():
                raise ValueError("No se encontró la plantilla Word o no se ha establecido")
            
            if not self.output_folder:
                raise ValueError("No se ha seleccionado una carpeta de salida para los boletines")
                
            # Asegurarse que la carpeta existe
            self.output_folder.mkdir(exist_ok=True)
            
            # 70% - Preparar archivo
            if callback:
                callback("📝 Preparando archivo del boletín (70%)")
            # Crear nombre de archivo seguro
            safe_name = nombre_completo.replace("/", "-").replace("\\", "-")
            base_filename = f"{student_id} - {safe_name}"
            
            # Crear paths para Word y PDF
            output_path_word = self.output_folder / f"{base_filename}.docx"
            output_path_pdf = self.output_folder / f"{base_filename}.pdf"
            
            # Asegurar que las rutas sean absolutas
            output_path_word = output_path_word.resolve()
            output_path_pdf = output_path_pdf.resolve()
            
            # Crear una copia de la plantilla
            import shutil
            shutil.copy2(self.word_template, output_path_word)
            
            # 75% - Calcular estadísticas académicas
            if callback:
                callback("📊 Calculando estadísticas académicas (75%)")
            
            # Debug - guardar datos en archivo
            with open('debug_notas.txt', 'w') as f:
                f.write(f"Datos del estudiante {student_id}:\n")
                for col, val in student_row.items():
                    f.write(f"{col}: {val}\n")
            
            promedio, materias_perdidas = self.calculate_academic_stats(student_row)
            
            # 80% - Preparar datos
            if callback:
                callback("✍ Preparando datos del estudiante (80%)")
            replacements = {}
            for field, value in {
                'nombre_completo': nombre_completo,
                'id': student_id,
                'grado': str(grado),
                'grupo': str(grupo) if grupo else 'N/A',
                'periodo': str(periodo) if periodo else 'N/A',
                'materias_perdidas': str(materias_perdidas),
                'promedio': f"{promedio:.2f}"
            }.items():
                for variant in self.word_fields_mapping[field]:
                    replacements[variant] = value
            
            # 90% - Procesar documento
            if callback:
                callback("📄 Generando boletín (90%)")
            # If the grade is 8, exclude the 'Investigación' area from mapping
            try:
                grado_int = int(str(grado).strip())
            except Exception:
                grado_int = None

            areas_mapping_to_use = self.areas_mapping
            if grado_int == 8:
                # Create a shallow copy and remove 'Investigación' if present
                areas_mapping_to_use = {k: v for k, v in self.areas_mapping.items() if k != 'Investigación'}

            self.process_word_document(
                str(output_path_word),
                str(output_path_pdf),
                replacements,
                student_row,
                callback,
                areas_mapping=areas_mapping_to_use
            )
            
            # 100% - Finalizar
            if callback:
                callback("✅ Boletín completado (100%)")
            
            return {
                'word': str(output_path_word),
                'pdf': str(output_path_pdf)
            }
            
        except Exception as e:
            if callback:
                callback(f"❌ Error: {str(e)}")
            raise

    def process_all_students(self, periodo, callback=None):
        """Procesa todos los estudiantes para un periodo dado"""
        if not self.excel_path:
            raise ValueError("No se ha seleccionado un archivo Excel")
            
        # Iterar sobre todas las hojas (grados)
        try:
            xlsx = pd.ExcelFile(self.excel_path)
            total_sheets = len(xlsx.sheet_names)
            
            for i, grado in enumerate(xlsx.sheet_names):
                if grado.lower().strip() == 'consolidado':
                    continue
                    
                if callback:
                    callback(f"📚 Procesando grado: {grado}")
                
                df = pd.read_excel(self.excel_path, sheet_name=grado)
                df = self.normalize_column_names(df)
                total_students = len(df)
                
                for j, row in df.iterrows():
                    student_field = str(row['estudiante']).strip()
                    if student_field.lower() == 'nan' or student_field == '':
                        continue
                        
                    student_id, student_name, _, _ = self.parse_student_info(student_field)
                    if student_id:
                        if callback:
                            callback(f"📝 Procesando estudiante: {student_name}")
                        
                        grupo = str(row.get('GRUPO', '')).strip()
                        if pd.isna(grupo) or grupo.lower() == 'nan' or grupo == '':
                            grupo = 'N/A'
                            
                        self.process_student(student_id, grado, grupo, periodo, callback=callback)
        except Exception as e:
            if callback:
                callback(f"❌ Error: {str(e)}")
            raise

if __name__ == "__main__":
    # Iniciar la interfaz gráfica
    gui = NotadorGUI()
    gui.run()
